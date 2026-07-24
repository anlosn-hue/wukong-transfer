#!/usr/bin/env python3
"""
声誉风险应急预案 / 回撤预案生成器
用法：python gen_plan_doc.py --type activity|system --output <docx路径> --data <json路径>

data JSON 字段（活动类 activity）：
    complaint_handling  str   投诉舆情处理段落
    customer_qa         list  [[情景状态, 场景说明, 回答, 注意事项, 备注], ...]
    media_qa            list  [[媒体问题, 回答, 适用场景, 备注], ...]
    optimization        str   方案优化段落
    contact             str   联系人段落

data JSON 字段（系统类 system）：
    system_name         str   系统名称（用于标题和口径）
    complaint_handling  str   投诉舆情处理段落
    customer_qa         list  同上
    media_qa            list  同上
    optimization        str   方案优化段落
    contact_lines       list  [str, ...]  联系人多行（业务+技术各一行）
"""
import json
import argparse
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_table(doc, headers: list, rows: list):
    """添加带表头的表格，表头加粗。"""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for ri, row_data in enumerate(rows):
        tr = table.rows[1 + ri]
        for ci, val in enumerate(row_data):
            tr.cells[ci].text = str(val) if val else ''

    return table


def _activity_plan(doc, data: dict):
    title_para = doc.add_heading('应急预案', 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1.  投诉、舆情处理', 1)
    doc.add_paragraph(data.get('complaint_handling',
        '发生客户投诉、舆情问题时，及时安抚处理，若发生严重客户投诉及舆情问题时，'
        '结合相关情况，停止投放相关活动及策略。'))

    doc.add_heading('2.  备答口径', 1)

    doc.add_heading('（一）客服人员应对客户咨询/投诉备答口径', 2)
    _add_table(doc,
               headers=['情景状态', '场景说明', '回 答', '注意事项', '备注'],
               rows=data.get('customer_qa', []))

    doc.add_heading('（二）媒体应答口径', 2)
    _add_table(doc,
               headers=['媒体问题', '回答', '适用场景', '备注'],
               rows=data.get('media_qa', []))

    doc.add_heading('3.  方案优化', 1)
    doc.add_paragraph(data.get('optimization',
        '根据投诉、舆情相关内容，优化策略投放逻辑及相关物料。'))

    doc.add_heading('4.  联系人', 1)
    doc.add_paragraph(data.get('contact',
        '针对活动如遇到客户投诉、抱怨问题，请联系数字运营部对口联系人（请勿提供给客户）。\n'
        '一旦发生投诉，力争第一时间解决。'))

    doc.add_paragraph()
    doc.add_paragraph('【具体日期，以公告发布的前置流程最终审批通过为准】')


def _system_plan(doc, data: dict):
    system_name = data.get('system_name', '系统')
    title_para = doc.add_heading('回撤预案', 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1.  投诉、舆情处理', 1)
    doc.add_paragraph(data.get('complaint_handling',
        f'发生客户投诉、舆情问题时，及时安抚处理，若发生严重客户投诉及舆情问题时，'
        f'结合相关情况，协调相关供应商停止切换或立即进行回切。'))

    doc.add_heading('2.  备答口径', 1)

    doc.add_heading('（一）客服人员应对客户咨询/投诉备答口径', 2)
    _add_table(doc,
               headers=['情景状态', '场景说明', '回 答', '注意事项', '备注'],
               rows=data.get('customer_qa', []))

    doc.add_heading('（二）媒体应答口径', 2)
    _add_table(doc,
               headers=['媒体问题', '回答', '适用场景', '备注'],
               rows=data.get('media_qa', []))

    doc.add_heading('3.  方案优化', 1)
    doc.add_paragraph(data.get('optimization',
        f'若发生严重客户投诉及舆情问题时，结合相关情况，协调合作供应商停止切换或立即进行回切。'))

    doc.add_heading('4.  联系人', 1)
    for line in data.get('contact_lines', [
        '业务联系人：【待填写】（请勿提供给客户）',
        '技术联系人：【待填写】（请勿提供给客户）',
    ]):
        doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_paragraph('一旦发生投诉，力争第一时间解决。')
    doc.add_paragraph()
    doc.add_paragraph('【具体日期，以公告发布的前置流程最终审批通过为准】')


def generate_plan_doc(plan_type: str, output_path: str, data: dict) -> None:
    doc = Document()

    if plan_type == 'activity':
        _activity_plan(doc, data)
    elif plan_type == 'system':
        _system_plan(doc, data)
    else:
        raise ValueError(f'plan_type 必须为 activity 或 system，收到：{plan_type!r}')

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ 已生成预案：{output_path}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='生成声誉风险应急/回撤预案')
    parser.add_argument('--type',   required=True, choices=['activity', 'system'])
    parser.add_argument('--output', required=True, help='输出 docx 路径')
    parser.add_argument('--data',   required=True, help='数据 JSON 文件路径')
    args = parser.parse_args()

    with open(args.data, encoding='utf-8') as f:
        data = json.load(f)

    generate_plan_doc(args.type, args.output, data)
