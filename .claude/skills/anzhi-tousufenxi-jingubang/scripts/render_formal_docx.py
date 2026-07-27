# -*- coding: utf-8 -*-
"""正式报告 docx 渲染：封面+目录+彩色callout框+表格+matplotlib图表。
用法：python render_formal_docx.py <报告月份dir> <config.yaml>"""
import io, json, sys
from pathlib import Path
import yaml
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import chart_images
import docx_footnotes as fn
import report_notes
import report_outline

BLUE = RGBColor(0x1a, 0x4d, 0x8f)
AMBER = RGBColor(0xa3, 0x69, 0x0f)
RED = RGBColor(0xa3, 0x23, 0x1b)

def _add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-1" \h \z \u'
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "（首次打开请右键此处选择「更新域」以生成目录页码）"
    fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "end")
    r = run._r
    for el in (fld1, instr, fld2, placeholder, fld3):
        r.append(el)

def _add_cover(doc, month):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("总行数字运营部"); r.font.size = Pt(22); r.font.bold = True
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("客户投诉分析报告"); r2.font.size = Pt(18)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 期间标签按分析周期渲染（月/季/半年/年度通用），不写死"数据截至X月末"
    r3 = p3.add_run(f"（{report_notes.period_label(month)}）")
    r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

def _set_cell_shading(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def _set_cell_left_border(cell, hex_color, sz=32):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:color"), hex_color)
    borders.append(left)
    tcPr.append(borders)

def _add_callout(doc, label, text, border_color, label_color, fill_hex):
    """彩色边框callout框：单格表格模拟盒子（python-docx无段落边框/底纹高层API，需操作tcPr原始OOXML）。
    border_color/label_color 传 RGBColor（取其hex用于边框），fill_hex 传"EEF4FB"式浅色底纹hex字符串。"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, fill_hex)
    _set_cell_left_border(cell, str(border_color))
    p = cell.paragraphs[0]
    r = p.add_run(label); r.font.bold = True; r.font.color.rgb = label_color
    cell.add_paragraph(text)
    return table

def _write(container, text, ctx, style=None):
    """写一个可能带脚注标记的段落。container 可以是 Document 也可以是表格单元格。"""
    if ctx is None:
        return container.add_paragraph(text)
    return fn.add_text_with_footnotes(container, text, ctx["fn"], ctx["notes"], style=style)


def _render_table_md(doc, table_md, ctx=None):
    """把 report_outline 产出的简易markdown表格（可能含多段，段落间用\n\n分隔）逐个渲染成docx表格；
    非表格的纯文本段落原样加成段落，"###"开头的当作小节标题。"""
    for block in table_md.split("\n\n"):
        lines = [l for l in block.splitlines() if l.strip()]
        table_lines = [l for l in lines if l.strip().startswith("|")]
        if not table_lines or len(table_lines) < 2:
            if block.strip().startswith("###"):
                doc.add_heading(block.strip().lstrip("#").strip(), level=2)
            elif any(l.strip().startswith("- ") for l in lines):
                # markdown 列表逐条成段：整块直写会把多条挤成一段，还把 '-' 原样印进公文
                for l in lines:
                    s = l.strip()
                    if s.startswith("- "):
                        _write(doc, s[2:].strip(), ctx, style="List Bullet")
                    elif s:
                        _write(doc, s, ctx)
            else:
                # 多行普通文本逐行成段：整块直写会把几段挤成一段（叙述字段常见）
                for l in lines:
                    if l.strip():
                        _write(doc, l.strip(), ctx)
            continue
        rows = [table_lines[0]] + table_lines[2:]  # 第2行固定是 |---|---| 分隔行，跳过它，不做字符集判断
        # 表格单元格内不落脚注（Word 表内脚注排版不稳），标记直接去掉
        cells = lambda l: [fn.strip_markers(c.strip()) for c in l.strip().strip("|").split("|")]
        head_cells = cells(rows[0])
        t = doc.add_table(rows=1, cols=len(head_cells))
        t.style = "Light Grid Accent 1"
        for c, text in zip(t.rows[0].cells, head_cells):
            c.text = text
        for r in rows[1:]:
            row_cells = t.add_row().cells
            for c, text in zip(row_cells, cells(r)):
                c.text = text

def _render_chart(doc, chart):
    if chart["type"] == "bar":
        png = chart_images.bar_chart_png(chart["labels"], chart["series"], chart["title"],
                                          xlabel=chart.get("xlabel"))
    else:
        png = chart_images.line_chart_png(chart["series"], chart["title"])
    if not png:
        return
    doc.add_picture(io.BytesIO(png), width=Inches(6))

def _render_note(doc, note):
    if not note:
        return
    for line in note.split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def run(report_dir, config):
    report_dir = Path(report_dir)
    sw = config.get("产出物开关", {})
    if not sw.get("docx", True):
        return
    metrics = json.loads((report_dir / "指标.json").read_text(encoding="utf-8"))
    np_ = report_dir / "叙述.json"
    narrative = json.loads(np_.read_text(encoding="utf-8")) if np_.exists() else {}
    sp = report_dir / "摘要.json"
    summaries = json.loads(sp.read_text(encoding="utf-8")) if sp.exists() else {}

    doc = Document()
    ctx = {"fn": fn.FootnoteManager(), "notes": report_notes.footnotes(metrics)}
    _add_cover(doc, metrics["月份"])
    toc_label = doc.add_paragraph()
    tr = toc_label.add_run("目录"); tr.font.size = Pt(16); tr.font.bold = True
    # 注意：目录标签用普通段落而非 Heading 1，避免它自己也被 TOC \o "1-1" 域扫描进目录列表（自我引用）
    _add_toc_field(doc)
    doc.add_page_break()

    CN_NUM = ["一", "二", "三", "四", "五", "六", "七", "八",
              "九", "十", "十一", "十二", "十三", "十四"]
    no = iter(CN_NUM)

    # 首章「报告说明」：数据范围/预警规则/模型清单/分析层级，全文脚注的定义都在这一章
    intro = report_notes.build_intro(metrics)
    doc.add_heading(f"{next(no)}、{intro['title']}", level=1)
    _render_table_md(doc, intro["table_md"], ctx)

    doc.add_heading(f"{next(no)}、结论摘要", level=1)
    findings = narrative.get("关键发现速览", [])
    if findings:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.rows[0].cells[0]
        _set_cell_shading(cell, "FDF6EA")
        _set_cell_left_border(cell, str(AMBER))
        hp = cell.paragraphs[0]
        hr = hp.add_run("★ 关键发现速览"); hr.font.bold = True; hr.font.color.rgb = AMBER
        for f in findings:
            fp = cell.add_paragraph(style="List Bullet")
            fr = fp.add_run(f"{f['标签']}："); fr.font.bold = True
            for chunk, key in fn.split_markers(f["文本"]):
                if chunk:
                    fp.add_run(chunk)
                if key and ctx["notes"].get(key):
                    ctx["fn"].add_ref(fp, ctx["notes"][key])
    for line in narrative.get("结论摘要", []):
        _write(doc, line, ctx)

    outline, appendix = report_outline.split_appendix(
        report_outline.build_outline(metrics, summaries, narrative))

    def emit(sec):
        doc.add_heading(f"{next(no)}、{sec['title']}", level=1)
        if sec["narrative"]:
            # 与 html 对齐：叙述支持分段/分点/加粗，不再当成单个段落
            _render_table_md(doc, sec["narrative"], ctx)
        if sec["insight"]:
            _add_callout(doc, "■ 数据洞察", fn.strip_markers(sec["insight"]), BLUE, BLUE, "EEF4FB")
        if sec["risk"]:
            _add_callout(doc, "⚠ 风险提示", fn.strip_markers(sec["risk"]), RED, RED, "FBECEB")
        _render_table_md(doc, sec["table_md"], ctx)
        if sec["chart"]:
            _render_chart(doc, sec["chart"])
        _render_note(doc, sec.get("note"))

    for sec in outline:
        emit(sec)

    doc.add_heading(f"{next(no)}、策略建议", level=1)
    for s in narrative.get("策略建议", []):
        p = doc.add_paragraph(); r = p.add_run(s["标题"]); r.font.bold = True; r.font.color.rgb = BLUE
        _write(doc, s["内容"], ctx)

    for sec in appendix:       # 附件清单固定排最后（用户 2026-07-25）
        emit(sec)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"生成：{metrics['生成时间']} · 总行数字运营部声誉风险管理智能体·悟空（金箍棒）")
    fr.font.size = Pt(9); fr.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    ctx["fn"].finalize(doc)  # 必须在 save 之前挂 footnotes.xml
    doc.save(str(report_dir / "月度分析报告.docx"))

if __name__ == "__main__":
    cfg = yaml.safe_load(Path(sys.argv[2]).read_text(encoding="utf-8"))
    run(sys.argv[1], cfg)
    print("docx渲染完成")
