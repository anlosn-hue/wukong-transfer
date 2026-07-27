# -*- coding: utf-8 -*-
"""Word 真脚注支持（页面底部、自动编号）。

python-docx 没有脚注 API，也不在默认模板里带 word/footnotes.xml，所以这里手工造：
1) 收集正文里出现的脚注文本，按首次出现顺序编号（id 从 1 起；-1/0 是 Word 规定的
   分隔符与续页分隔符条目，必须存在，否则 Word 打开时会提示文档需要修复）；
2) 正文处插入 <w:footnoteReference w:id="N"/> 的上标 run；
3) doc.save() 之前调用 finalize()，把 footnotes.xml 作为 OPC 部件挂到 document 上。
同一段脚注文本重复出现时复用同一编号，不重复排注。
"""
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from marker_utils import strip_tips
from marker_utils import split_markers as _split_markers
from marker_utils import strip_markers as _strip_markers


# 转出口（渲染器从这里取）。docx 侧一律**连带剥掉悬浮标记**：Word 没有悬停形态，
# 明细另作附件；不剥的话正式公文里会直接漏出「〔tip:...〕」源码。
def strip_markers(text):
    return strip_tips(_strip_markers(text))


def split_markers(text):
    return [(strip_tips(chunk), key) for chunk, key in _split_markers(text)]

CT_FOOTNOTES = ("application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.footnotes+xml")
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_HEAD = f'<w:footnotes xmlns:w="{NS_W}">'
_SEPARATORS = (
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr>'
    '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
    '<w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr>'
    '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
    '<w:r><w:continuationSeparator/></w:r></w:p></w:footnote>')


def _xml_escape(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


class FootnoteManager:
    """一份 docx 一个实例。add_ref(paragraph, text) 落脚注，finalize(doc) 写部件。"""

    def __init__(self):
        self._texts = []          # 按编号顺序（下标 0 → id 1）
        self._index = {}          # 文本 → id

    def add_ref(self, paragraph, text):
        """在段落末尾追加一个上标脚注引用，返回编号；同一条脚注只在首次出现处落注。

        每个 w:id 只能被引用一次——实测同一 id 被引用两次时 Word 会多排一个空白脚注，
        所以重复出现的规则不再重复排注（正文里同一术语通常只需在首次出现时解释）。
        """
        text = text.strip()
        if text in self._index:
            return None
        self._texts.append(text)
        fid = len(self._texts)
        self._index[text] = fid
        run = paragraph.add_run()
        run.font.superscript = True
        run.font.size = Pt(9)
        ref = OxmlElement("w:footnoteReference")
        ref.set(qn("w:id"), str(fid))
        run._r.append(ref)
        return fid

    def finalize(self, doc):
        """把收集到的脚注写成 word/footnotes.xml 并关联到文档；无脚注则什么都不做。"""
        if not self._texts:
            return
        body = []
        for i, text in enumerate(self._texts, 1):
            body.append(
                f'<w:footnote w:id="{i}"><w:p><w:pPr>'
                f'<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
                f'<w:r><w:rPr><w:vertAlign w:val="superscript"/><w:sz w:val="18"/></w:rPr>'
                f'<w:footnoteRef/></w:r>'
                f'<w:r><w:rPr><w:sz w:val="18"/></w:rPr>'
                f'<w:t xml:space="preserve"> {_xml_escape(text)}</w:t></w:r>'
                f'</w:p></w:footnote>')
        xml = (_HEAD + _SEPARATORS + "".join(body) + "</w:footnotes>").encode("utf-8")
        part = Part(PackURI("/word/footnotes.xml"), CT_FOOTNOTES, xml, doc.part.package)
        doc.part.relate_to(part, RT.FOOTNOTES)


def add_text_with_footnotes(doc, text, fnmgr, notes, style=None):
    """把含标记 〔fn:KEY〕 的文本写成一个段落，标记处落真脚注。

    notes: {KEY: 脚注全文}。未登记的 KEY 原样保留标记文本，以便渲染时一眼看出漏配。
    """
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for chunk, key in split_markers(text):
        if chunk:
            # markdown 的 **加粗** 在 docx 里要落成 bold run，否则星号会原样印进正式报告。
            # 落单的 ** 按字面处理——否则后半段会被整段加粗（同 html 侧 _escape_bold）
            if chunk.count("**") % 2:
                p.add_run(chunk)
            else:
                for i, seg in enumerate(chunk.split("**")):
                    if seg:
                        p.add_run(seg).bold = bool(i % 2)
        if key is not None:
            note = notes.get(key)
            if note:
                fnmgr.add_ref(p, note)
            else:
                p.add_run(f"〔fn:{key}〕")
    return p


