# -*- coding: utf-8 -*-
import json
from docx import Document
import render_formal_docx

METRICS = {"月份": "2026-05", "生成时间": "2026-07-07 09:00",
           "预警汇总": [{"级别": "红", "表": "督办", "问题点": "A/B", "依据": "比照命中",
                       "来源模型": "督办投诉比照"}],
           "模型": {"排名分析": {"指标": {"投诉": {"本月": []}, "督办": {"本月": []}}, "md": ""},
                    "督办投诉比照": {"指标": {}, "md": "### 督办×投诉比照\n\n无"},
                    "活动关联": {"指标": {}, "md": "### 活动\n\n无"},
                    "惯犯": {"指标": {}, "md": "### 惯犯\n\n无"}}}
NARRATIVE = {"关键发现速览": [{"标签": "红警", "文本": "A/B问题点触发红警"}],
             "结论摘要": ["1. 本月投诉整体平稳。"],
             "章节": {"投诉情况": {"叙述": "投诉平稳。", "洞察": None, "风险提示": None},
                     "督办情况": {"叙述": "督办平稳。", "洞察": None, "风险提示": None},
                     "督办转投诉预警": {"叙述": "A/B触发红警。", "洞察": None,
                                       "风险提示": "建议核实。"},
                     "在途活动关联": {"叙述": "无在途活动。", "洞察": None, "风险提示": None},
                     "重复发现问题清单": {"叙述": "数据积累中。", "洞察": None, "风险提示": None}},
             "策略建议": [{"标题": "建议一：加强核实", "内容": "对A/B问题点加强核实。"}]}

def _write_inputs(tmp_path):
    d = tmp_path / "报告" / "2026-05"; d.mkdir(parents=True)
    (d / "指标.json").write_text(json.dumps(METRICS, ensure_ascii=False), encoding="utf-8")
    (d / "叙述.json").write_text(json.dumps(NARRATIVE, ensure_ascii=False), encoding="utf-8")
    return d

def _full_text(doc):
    """doc.paragraphs 只含顶层段落，不含表格单元格内的段落（callout/关键发现速览已改为单格表格盒子）；
    这里补上表格内文本，得到完整的可见文字。"""
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)

def test_run_produces_docx_with_cover_and_toc_field(tmp_path):
    report_dir = _write_inputs(tmp_path)
    cfg = {"产出物开关": {"docx": True}}
    render_formal_docx.run(report_dir, cfg)
    out = report_dir / "月度分析报告.docx"
    assert out.exists()
    doc = Document(str(out))
    full_text = _full_text(doc)
    assert "总行数字运营部" in full_text and "客户投诉分析报告" in full_text
    assert "2026" in full_text
    xml = doc.element.xml
    assert "TOC" in xml and "fldChar" in xml  # 目录域已插入
    assert "关键发现速览" in full_text
    assert "建议一：加强核实" in full_text
    assert "总行数字运营部声誉风险管理智能体·悟空" in full_text

def test_run_skips_when_docx_switch_off(tmp_path):
    report_dir = _write_inputs(tmp_path)
    render_formal_docx.run(report_dir, {"产出物开关": {"docx": False}})
    assert not (report_dir / "月度分析报告.docx").exists()

def test_run_degrades_gracefully_without_narrative_file(tmp_path):
    d = tmp_path / "报告" / "2026-05"; d.mkdir(parents=True)
    (d / "指标.json").write_text(json.dumps(METRICS, ensure_ascii=False), encoding="utf-8")
    render_formal_docx.run(d, {"产出物开关": {"docx": True}})
    assert (d / "月度分析报告.docx").exists()  # 叙述.json缺失时仍能生成（章节叙述为空）

def test_run_renders_table_and_callout_and_chart(tmp_path):
    metrics = dict(METRICS)
    metrics["模型"] = dict(METRICS["模型"])
    metrics["模型"]["排名分析"] = {
        "指标": {"投诉": {"本月": [{"问题点": "借记卡限额", "笔数": 96, "占比": "39.0%", "变动": ""}]},
                "督办": {"本月": []}},
        "md": ""}
    d = tmp_path / "报告" / "2026-05"; d.mkdir(parents=True)
    (d / "指标.json").write_text(json.dumps(metrics, ensure_ascii=False), encoding="utf-8")
    (d / "叙述.json").write_text(json.dumps(NARRATIVE, ensure_ascii=False), encoding="utf-8")
    render_formal_docx.run(d, {"产出物开关": {"docx": True}})
    doc = Document(str(d / "月度分析报告.docx"))
    full_text = _full_text(doc)
    assert "建议核实" in full_text  # 风险提示callout文案出现
    tables_text = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "借记卡限额" in tables_text and "96" in tables_text  # 表格真正渲染为docx原生表格
    assert len(doc.inline_shapes) >= 1  # 至少一张图表图片被嵌入
    xml = doc.element.xml
    assert "w:shd" in xml and "w:tcBorders" in xml  # callout底纹+左边框OOXML确实写入

def test_render_table_md_skips_separator_row_by_position(tmp_path):
    metrics = dict(METRICS)
    metrics["模型"] = dict(METRICS["模型"])
    metrics["模型"]["排名分析"] = {
        "指标": {"投诉": {"本月": [{"问题点": "甲问题", "笔数": 10, "占比": "50.0%", "变动": ""},
                                  {"问题点": "乙问题", "笔数": 5, "占比": "25.0%", "变动": ""}]},
                "督办": {"本月": []}},
        "md": ""}
    d = tmp_path / "报告" / "2026-05"; d.mkdir(parents=True)
    (d / "指标.json").write_text(json.dumps(metrics, ensure_ascii=False), encoding="utf-8")
    (d / "叙述.json").write_text(json.dumps(NARRATIVE, ensure_ascii=False), encoding="utf-8")
    render_formal_docx.run(d, {"产出物开关": {"docx": True}})
    doc = Document(str(d / "月度分析报告.docx"))
    tousu_table = next(t for t in doc.tables if t.rows[0].cells[0].text == "#")
    # 表头行应是第0行、无分隔符行混入、数据行按顺序对应
    assert [c.text for c in tousu_table.rows[0].cells] == ["#", "诉点", "笔数", "占比"]
    assert [c.text for c in tousu_table.rows[1].cells] == ["1", "甲问题", "10", "50.0%"]
    assert [c.text for c in tousu_table.rows[2].cells] == ["2", "乙问题", "5", "25.0%"]


def test_docx_悬浮标记退化为标签不漏源码(tmp_path):
    """Word 没有悬停形态：〔tip:标签¦明细〕必须只留标签，明细走附件。
    正文段落、callout、表格单元格三条路径都要覆盖——任一条漏掉都会把源码印进公文。"""
    import copy
    d = tmp_path / "报告" / "2026-05"; d.mkdir(parents=True)
    n = copy.deepcopy(NARRATIVE)
    n["结论摘要"] = ["高敏感共〔tip:73案¦06-07 零售信贷部‖06-09 财富管理部〕。"]
    n["章节"]["督办转投诉预警"]["风险提示"] = "重点看〔tip:监管49案¦某条明细〕。"
    n["章节"]["兜底类目专项分析"] = {
        "叙述": "见下。", "洞察": None, "风险提示": None,
        "表格": "| 类别 | 说明 |\n|---|---|\n| 甲 | 〔tip:4案¦表内明细〕 |"}
    (d / "指标.json").write_text(json.dumps(METRICS, ensure_ascii=False), encoding="utf-8")
    (d / "叙述.json").write_text(json.dumps(n, ensure_ascii=False), encoding="utf-8")
    render_formal_docx.run(d, {})
    text = _full_text(Document(str(d / "月度分析报告.docx")))
    assert "〔tip:" not in text and "¦" not in text and "‖" not in text
    assert "高敏感共73案。" in text            # 正文段落
    assert "重点看监管49案。" in text          # callout
    assert "4案" in text and "表内明细" not in text  # 表格单元格


def test_docx_专项章节的markdown列表渲染成项目符号(tmp_path):
    """（八）高敏感条目改成分点后，docx 若不认 '- ' 会把整段挤成一个段落、
    并把 '-' 原样印出来。每条要各自成段并挂 List Bullet 样式。"""
    import copy
    d = tmp_path / "报告" / "2026-05"; d.mkdir(parents=True)
    n = copy.deepcopy(NARRATIVE)
    n["章节"]["兜底类目专项分析"] = {
        "叙述": "见下。", "洞察": None, "风险提示": None,
        "表格": "### （八）高敏感\n\n- **甲类 3案**：说明甲。\n- **乙类 4案**：说明乙。"}
    (d / "指标.json").write_text(json.dumps(METRICS, ensure_ascii=False), encoding="utf-8")
    (d / "叙述.json").write_text(json.dumps(n, ensure_ascii=False), encoding="utf-8")
    render_formal_docx.run(d, {})
    doc = Document(str(d / "月度分析报告.docx"))
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    texts = [p.text for p in bullets]
    assert "甲类 3案：说明甲。" in texts and "乙类 4案：说明乙。" in texts
    assert not any(t.lstrip().startswith("-") for t in texts)   # 不得留下 markdown 的 '-'
    assert "**" not in _full_text(doc)                          # 粗体落成 run，不是字面星号


def test_docx_单个星号对不成对时按字面处理():
    """同 html：落单的 ** 不得把后文全部加粗。"""
    from docx import Document as _D
    import docx_footnotes as fn
    doc = _D()
    fn.add_text_with_footnotes(doc, "甲**乙：说明", fn.FootnoteManager(), {})
    p = doc.paragraphs[-1]
    assert p.text == "甲**乙：说明"
    assert not any(r.bold for r in p.runs)


def test_docx_章节叙述支持分段与分点(tmp_path):
    """docx 侧同样：叙述里的换行要各自成段，'- ' 要成项目符号。"""
    import copy
    d = tmp_path / "报告" / "2026-05"; d.mkdir(parents=True)
    n = copy.deepcopy(NARRATIVE)
    n["章节"]["投诉情况"]["叙述"] = ("总述一句。\n- **第一位** 甲：1038条\n- **第二位** 乙：228条\n"
                                    "环比方面，另起一段。")
    (d / "指标.json").write_text(json.dumps(METRICS, ensure_ascii=False), encoding="utf-8")
    (d / "叙述.json").write_text(json.dumps(n, ensure_ascii=False), encoding="utf-8")
    render_formal_docx.run(d, {})
    doc = Document(str(d / "月度分析报告.docx"))
    texts = [p.text for p in doc.paragraphs]
    assert "总述一句。" in texts and "环比方面，另起一段。" in texts   # 各自成段
    bl = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert "第一位 甲：1038条" in bl and "第二位 乙：228条" in bl
    assert "**" not in _full_text(doc)
