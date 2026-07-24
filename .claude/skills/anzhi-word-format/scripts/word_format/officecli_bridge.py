# word_format/officecli_bridge.py
"""
officecli_bridge.py
word_format 与 officecli 的唯一交互点。

分工：word_format 用自己的标题识别规则给段落打 w:outlineLvl（officecli 不知道
本单位公文的标题编号长什么样）；TOC 域的插入与页码刷新委托给 officecli。

实测依据（2026-07-19，officecli 1.0.138）：标题段落若没有 outlineLvl，
officecli 生成的 TOC 域刷新后是「未找到目录项」；打上之后目录与页码均正确。

降级策略：officecli 缺失/报错/超时一律只写 report warning 并返回 False，
绝不抛异常——TOC 是可选增强，不应让整个格式化流程失败。
"""
from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from word_format.config_types import BodyConfig
from word_format.report_types import FormatReport

TOC_LEVELS = (1, 2, 3)          # 目录只收一~三级
_EXE = "officecli"

# ECMA-376 CT_PPr 子元素顺序中，排在 outlineLvl 之后的元素。
# outlineLvl 必须插在这些元素中第一个出现者之前。
_AFTER_OUTLINE_LVL = ("w:rPr",)


def mark_outline_levels(doc: Document, config: BodyConfig,
                        report: FormatReport) -> int | None:
    """给一~三级标题段落打 w:outlineLvl（0-based），供 officecli 收录目录项。

    用直接段落级覆盖，不依赖 Word 内置 Heading 命名样式——本单位公文正文
    不使用 Heading 样式，靠编号正则识别标题。

    **返回首个标题的 0-based 段落下标**（无标题时返回 None）。这个值是 TOC
    的插入锚点：officecli add 默认 append 到文末，必须显式告诉它插哪。

    同时给首个标题段落打 w:pageBreakBefore，让目录独占一页（公文惯例）。
    """
    patterns = [(h.level, re.compile(h.pattern)) for h in config.headings]
    n_marked = 0
    first_idx: int | None = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        level = _detect_level(text, patterns)
        if level not in TOC_LEVELS:
            continue
        _set_outline_lvl(para, level - 1)
        if first_idx is None:
            first_idx = i
        n_marked += 1

    if first_idx is not None:
        _set_page_break_before(doc.paragraphs[first_idx])   # 目录独占一页
    if n_marked:
        report.add("style", "applied", f"已标记 {n_marked} 处标题大纲级别（供目录收录）")
    else:
        report.add("style", "warning", "未识别到一~三级标题，目录将为空，请检查文档编号格式")
    return first_idx


def _set_page_break_before(para) -> None:
    """给段落加 w:pageBreakBefore，遵守 CT_PPr 子元素顺序。

    顺序约束：pageBreakBefore 排在 outlineLvl 与 rPr **之前**。
    """
    pPr = para._p.get_or_add_pPr()
    if pPr.find(qn("w:pageBreakBefore")) is not None:
        return
    el = OxmlElement("w:pageBreakBefore")
    anchor = None
    for tag in ("w:outlineLvl", "w:rPr"):
        found = pPr.find(qn(tag))
        if found is not None:
            anchor = found
            break
    if anchor is not None:
        anchor.addprevious(el)
    else:
        pPr.append(el)


def _detect_level(text: str, patterns: list[tuple[int, re.Pattern]]) -> int:
    for level, pattern in patterns:
        if pattern.match(text):
            return level
    return 0


def _set_outline_lvl(para, val: int) -> None:
    """写入 w:outlineLvl，遵守 CT_PPr 子元素顺序（必须在 w:rPr 之前）。"""
    pPr = para._p.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        anchor = None
        for tag in _AFTER_OUTLINE_LVL:
            found = pPr.find(qn(tag))
            if found is not None:
                anchor = found
                break
        if anchor is not None:
            anchor.addprevious(outline)     # 插到 rPr 之前，保证 schema 合法
        else:
            pPr.append(outline)
    outline.set(qn("w:val"), str(val))


def insert_toc(docx_path: Path, report: FormatReport,
               anchor_idx: int | None = None) -> bool:
    """调 officecli 插入 TOC 域。失败一律降级，返回 False。

    **anchor_idx 是 mark_outline_levels 返回的首个标题 0-based 下标。**
    officecli 的 `add` 默认 append 到文档**末尾**（help 原文：
    "If omitted, appends to end"），不传位置会得到一份目录长在落款后面的文档。

    索引换算（已实测验证）：python-docx 的 0-based `doc.paragraphs[i]`
    等于 officecli 的 1-based `/body/p[i+1]`。要把目录插在首个标题**之前**，
    就是插在它前一段**之后** → `--after /body/p[anchor_idx]`。
    anchor_idx == 0（文档以标题开头、无标题块）时改用 `--index 0`。
    """
    argv = ["add", str(docx_path), "/", "--type", "toc"]
    if anchor_idx is None:
        report.add("style", "warning", "无标题锚点，目录将追加到文档末尾")
    elif anchor_idx == 0:
        argv += ["--index", "0"]
    else:
        argv += ["--after", f"/body/p[{anchor_idx}]"]
    argv += ["--prop", "levels=1-3",
             "--prop", "title=目录",
             "--prop", "hyperlinks=true",
             "--prop", "pagenumbers=true"]
    return _run_officecli(
        argv,
        timeout=30,
        ok_detail="已插入目录（officecli add --type toc，收录一~三级标题）",
        fail_prefix="目录插入失败",
        report=report,
    )


def refresh_toc(docx_path: Path, report: FormatReport) -> bool:
    """调 officecli 刷新域，把目录页码固化为真实值。失败一律降级，返回 False。"""
    return _run_officecli(
        ["refresh", str(docx_path)],
        timeout=60,
        ok_detail="已刷新域（目录页码已固化）",
        fail_prefix="域刷新失败",
        report=report,
        parse_backend=True,
    )


def close_document(docx_path: Path, report: FormatReport) -> bool:
    """结束 officecli 的常驻进程，释放文件句柄。

    **必须调用。** officecli 会起常驻进程把文档留在内存里（help 原文：
    close = "Flush in-memory changes to disk and stop the resident (releases
    the file)"）。不 close 就用 python-docx 打开写入会抛
    `PermissionError: [Errno 13] Permission denied`（实测踩到过）。
    """
    return _run_officecli(
        ["close", str(docx_path)],
        timeout=30,
        ok_detail="已释放 officecli 文件句柄",
        fail_prefix="释放文件句柄失败",
        report=report,
    )


# TOC 标题样式在 OOXML 里的 w:name（styleId 会被本地化成 "TOC"/数字，不可靠）
_TOC_HEADING_STYLE_NAMES = ("TOC Heading",)


def blacken_toc_heading(docx_path: Path, report: FormatReport) -> bool:
    """把 officecli 插入的目录标题「目录」改为纯黑。

    实测：该样式 styleId 为 `TOC`（w:name = "TOC Heading"），自身**不带**
    w:color，蓝色是从 `basedOn` 的标题样式继承来的（中文 Word 文档里标题
    styleId 被本地化成 "1"/"2"，不是 "Heading1"）。所以不去追继承链，
    直接在 TOC 样式上写死黑色覆盖继承值。

    必须在 close_document 之后调用（否则文件被占用）。
    """
    try:
        doc = Document(str(docx_path))
        n = 0
        for style_el in doc.styles.element.findall(qn("w:style")):
            name_el = style_el.find(qn("w:name"))
            if name_el is None or name_el.get(qn("w:val")) not in _TOC_HEADING_STYLE_NAMES:
                continue
            rPr = style_el.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                style_el.append(rPr)
            color_el = rPr.find(qn("w:color"))
            if color_el is None:
                color_el = OxmlElement("w:color")
                rPr.append(color_el)
            for attr in list(color_el.attrib):
                del color_el.attrib[attr]
            color_el.set(qn("w:val"), "000000")
            n += 1
        if n:
            doc.save(str(docx_path))
            report.add("style", "fix", "目录标题已改为纯黑")
        return True
    except Exception as exc:                       # noqa: BLE001 — 降级不中断
        report.add("style", "warning", f"目录标题改黑失败：{exc}（不影响目录内容）")
        return False


def _run_officecli(argv: list[str], timeout: int, ok_detail: str,
                   fail_prefix: str, report: FormatReport,
                   parse_backend: bool = False) -> bool:
    manual_hint = "请在 Word 中按 Ctrl+A 后按 F9 手动更新域"
    if shutil.which(_EXE) is None:
        report.add("style", "warning",
                   f"未检测到 officecli，目录/页码未生成。请安装 officecli 后重试，"
                   f"或{manual_hint}")
        return False
    try:
        proc = subprocess.run(
            [_EXE, *argv],
            capture_output=True, text=True, encoding="utf-8",
            timeout=timeout,
        )
    except subprocess.TimeoutExpired:
        report.add("style", "warning", f"{fail_prefix}：officecli 超时（>{timeout}s），{manual_hint}")
        return False
    except Exception as exc:                       # noqa: BLE001 — 降级不区分异常类型
        report.add("style", "warning", f"{fail_prefix}：{exc}，{manual_hint}")
        return False

    if proc.returncode != 0:
        err = (proc.stderr or proc.stdout or "").strip().splitlines()
        summary = err[0][:200] if err else f"退出码 {proc.returncode}"
        report.add("style", "warning", f"{fail_prefix}：{summary}，{manual_hint}")
        return False

    detail = ok_detail
    if parse_backend:
        m = re.search(r"backend:\s*(\w+)", proc.stdout or "")
        if m:
            backend = m.group(1)
            detail = f"{detail}（backend: {backend}）"
            if backend != "word":
                detail = f"{detail}——非 Word 后端，页码可能不准，建议核对"
    report.add("style", "applied", detail)
    return True
