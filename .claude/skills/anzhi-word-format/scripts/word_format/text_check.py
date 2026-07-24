# word_format/text_check.py
"""
text_check.py
文本检查与修复。所有规则从 TextCheckConfig 读取。
"""
from __future__ import annotations

import re
from typing import Any
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from word_format.config_types import TextCheckConfig, CustomFix, LeaderConfig, OrgConfig
from word_format.report_types import FormatReport


def run_checks(doc: Document, config: TextCheckConfig, report: FormatReport) -> None:
    _check_paragraphs(doc, config, report)
    if config.fix_quote_direction:
        _fix_chinese_quote_direction(doc, report)
    if config.fix_list:
        _fix_auto_numbering(doc, report)
    if config.custom_fixes:
        _run_custom_fixes(doc, config.custom_fixes, report)
    if config.leader_check_enabled:
        _run_leader_check(doc, config.leaders, report)
    if config.org_check_enabled:
        _run_org_check(doc, config.orgs, report)


def _check_paragraphs(doc: Document, config: TextCheckConfig, report: FormatReport) -> None:
    body_paras = list(doc.paragraphs)
    table_paras: list = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_paras.extend(cell.paragraphs)
    for para in body_paras:
        if config.fix_bullet_chars:
            _fix_bullet_chars_in_para(para, report)
        _run_char_fixes(para, config, report)
    # 表格单元格不做行首符号清理：开头的 - 是负数、·/- 是层级前缀，删了会改数据含义
    for para in table_paras:
        _run_char_fixes(para, config, report)


def _run_char_fixes(para: Any, config: TextCheckConfig, report: FormatReport) -> None:
    if config.fix_date_padding:
        _fix_date_in_para(para, report)
    if config.fix_zhizhi:
        _fix_jiezhi_in_para(para, report)
    if config.fix_quotes:
        _fix_quotes_in_para(para, report)


# ── 行首符号删除 ───────────────────────────────────────────────────────────────
# 符号后紧跟数字（如 -0.03、-14.45%）视为负数，不当作项目符号
_BULLET_CHAR_RE = re.compile(r"^[·\-](?!\s*[\d.])\s*")


def _fix_bullet_chars_in_para(para: Any, report: FormatReport) -> None:
    for run in para.runs:
        if not run.text:
            continue
        fixed, n = _BULLET_CHAR_RE.subn("", run.text)
        if n > 0:
            report.add("text", "fix",
                       f'行首符号删除："{para.text[:40].strip()}" 删除 "{run.text[0]}"')
            run.text = fixed
        break


# ── 日期修复 ──────────────────────────────────────────────────────────────────

_DATE_RE = re.compile(r"(\d{4}年)0*([1-9]\d*)月0*([1-9]\d*)日")


def _fix_date_in_para(para: Any, report: FormatReport) -> None:
    for run in para.runs:
        original = run.text
        fixed, n = _DATE_RE.subn(r"\g<1>\2月\3日", original)
        if n > 0:
            run.text = fixed
            report.add("text", "fix", f'日期修复："{original.strip()}" → "{fixed.strip()}"')


# ── 截止 → 截至 ───────────────────────────────────────────────────────────────

_JIEZHI_RE = re.compile(r"截止(?!到)")


def _fix_jiezhi_in_para(para: Any, report: FormatReport) -> None:
    for run in para.runs:
        original = run.text
        fixed, n = _JIEZHI_RE.subn("截至", original)
        if n > 0:
            run.text = fixed
            report.add("text", "fix", f'截至修复："{original.strip()}" → "{fixed.strip()}"')


# ── 英文直双引号 → 中文弯引号 ─────────────────────────────────────────────────

_OPEN_DQUOTE  = "“"   # 中文左双引号 “
_CLOSE_DQUOTE = "”"   # 中文右双引号 ”


def _fix_quotes_in_para(para: Any, report: FormatReport) -> None:
    """段落内直双引号 (") 按出现顺序交替转中文前/后引号，开合状态跨 run 持续。"""
    open_state = False
    n_fixed = 0
    for run in para.runs:
        if '"' not in run.text:
            continue
        chars = []
        for ch in run.text:
            if ch == '"':
                chars.append(_OPEN_DQUOTE if not open_state else _CLOSE_DQUOTE)
                open_state = not open_state
                n_fixed += 1
            else:
                chars.append(ch)
        run.text = "".join(chars)
    if n_fixed:
        report.add("text", "fix",
                   f'直引号转中文引号：{n_fixed} 处（段落："{para.text[:30].strip()}"）')


# ── 中文弯引号方向重排（pandoc smart quotes 误判修复）─────────────────────────

_OPEN_SQUOTE  = "‘"   # 中文左单引号 ‘
_CLOSE_SQUOTE = "’"   # 中文右单引号 ’
_DQUOTE_CHARS = (_OPEN_DQUOTE, _CLOSE_DQUOTE)
_SQUOTE_CHARS = (_OPEN_SQUOTE, _CLOSE_SQUOTE)


def _fix_chinese_quote_direction(doc: Document, report: FormatReport) -> None:
    """逐段落重排中文弯引号方向（双引号、单引号各自独立计数）。

    pandoc 的 smart quotes 靠空格边界判断开合，中文无空格时常判反；
    这里不管当前字符是开是闭，只按段内出现顺序重新交替赋值，纠正 pandoc 的误判。

    与既有 _fix_quotes_in_para 的分工：那个负责「英文直引号 → 中文弯引号」的
    字符替换，本函数负责「已是中文弯引号但方向错」的重排，跑在其之后。
    两者都是段落级，心智模型一致。

    为什么是段落级：对每段引号成对的正常文档，段落级与文档级结果完全相同；
    只在文档本身畸形（某段引号未闭合）时才分叉，此时段落级把损害关在那一段内，
    不会翻转其后本来正确的段落。段落级严格占优。
    """
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    _reassign_quote_pairs(all_paras, _DQUOTE_CHARS, _OPEN_DQUOTE, _CLOSE_DQUOTE,
                          "双引号", report, guard_apostrophe=False)
    _reassign_quote_pairs(all_paras, _SQUOTE_CHARS, _OPEN_SQUOTE, _CLOSE_SQUOTE,
                          "单引号", report, guard_apostrophe=True)


def _is_apostrophe(text: str, i: int) -> bool:
    """判定 text[i] 处的 U+2019 是否为英文撇号（don't / it's）而非右单引号。

    规则：两侧均为 ASCII 字母或数字。pandoc 对英文撇号正是输出 U+2019，
    若不加此护栏会被当成引号翻成左单引号，属数据损坏。
    """
    if text[i] != _CLOSE_SQUOTE:
        return False
    prev_ch = text[i - 1] if i > 0 else ""
    next_ch = text[i + 1] if i + 1 < len(text) else ""
    return (prev_ch.isascii() and prev_ch.isalnum()
            and next_ch.isascii() and next_ch.isalnum())


def _reassign_in_paragraph(para: Any, quote_chars: tuple[str, str],
                           open_ch: str, close_ch: str,
                           guard_apostrophe: bool) -> tuple[int, int]:
    """段落内交替赋开/闭。返回 (改动数, 参与计数的引号数)。

    先按 run 拼出全段文本用于撇号判定（撇号要看左右邻字符，可能跨 run），
    再按同一位置索引走 run 改写，保证位置对齐。
    """
    full = "".join(r.text for r in para.runs)
    skip: set[int] = set()
    if guard_apostrophe:
        skip = {i for i in range(len(full)) if _is_apostrophe(full, i)}

    idx = 0          # 段内交替计数
    pos = 0          # 段内全局字符位置（与 full 对齐）
    n_changed = 0
    n_quotes = 0
    for run in para.runs:
        if not any(ch in quote_chars for ch in run.text):
            pos += len(run.text)
            continue
        chars = []
        for ch in run.text:
            if ch in quote_chars and pos not in skip:
                new_ch = open_ch if idx % 2 == 0 else close_ch
                idx += 1
                n_quotes += 1
                if new_ch != ch:
                    n_changed += 1
                chars.append(new_ch)
            else:
                chars.append(ch)
            pos += 1
        run.text = "".join(chars)
    return n_changed, n_quotes


def _reassign_quote_pairs(paragraphs: list, quote_chars: tuple[str, str],
                          open_ch: str, close_ch: str, label: str,
                          report: FormatReport,
                          guard_apostrophe: bool = False) -> None:
    """逐段落对一类引号交替赋开/闭。

    某段为奇数时仍然修复（尽力修），并逐段记录警告供人工核对
    （取代早期「整类跳过」方案）。
    """
    total_changed = 0
    odd_paras: list[str] = []
    for para in paragraphs:
        changed, count = _reassign_in_paragraph(
            para, quote_chars, open_ch, close_ch, guard_apostrophe)
        total_changed += changed
        if count % 2 != 0:
            odd_paras.append(para.text)

    if total_changed:
        report.add("text", "fix",
                   f"{label}方向修正：{total_changed} 处按段内顺序重新交替赋值")
    for text in odd_paras:
        report.add("text", "warning",
                   f"{label}数为奇数，该段有落单引号，已尽力修复，请人工核对："
                   f"\"{text[:40].strip()}\"")


# ── custom_fixes ──────────────────────────────────────────────────────────────

def _run_custom_fixes(doc: Document, fixes: list[CustomFix], report: FormatReport) -> None:
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    for para in all_paras:
        for fix in fixes:
            for run in para.runs:
                if fix.source in run.text:
                    run.text = run.text.replace(fix.source, fix.target)
                    report.add("text", "fix",
                               f'自定义替换："{fix.source}" → "{fix.target}"'
                               f'（段落：{para.text[:30].strip()}）')


# ── leader_check ─────────────────────────────────────────────────────────────

def _run_leader_check(doc: Document, leaders: list[LeaderConfig],
                      report: FormatReport) -> None:
    if not leaders:
        return
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    full_text = "\n".join(p.text for p in all_paras)
    first_pos: dict[str, int] = {}
    for ldr in leaders:
        pos = full_text.find(ldr.name)
        if pos >= 0:
            first_pos[ldr.name] = pos

    mentioned = sorted(
        [(name, pos) for name, pos in first_pos.items()],
        key=lambda x: x[1]
    )
    leader_map = {ldr.name: ldr for ldr in leaders}
    for i in range(len(mentioned)):
        for j in range(i + 1, len(mentioned)):
            name_i, _ = mentioned[i]
            name_j, _ = mentioned[j]
            ldr_i = leader_map.get(name_i)
            ldr_j = leader_map.get(name_j)
            if ldr_i and ldr_j and ldr_i.rank > ldr_j.rank:
                report.add("leader", "warning",
                           f'领导人顺序错误：{name_i}（rank={ldr_i.rank}）'
                           f'出现在 {name_j}（rank={ldr_j.rank}）之前')


# ── org_check ─────────────────────────────────────────────────────────────────

def _run_org_check(doc: Document, orgs: list[OrgConfig], report: FormatReport) -> None:
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    full_text = "\n".join(p.text for p in all_paras)
    for org in orgs:
        if not org.abbr:
            continue
        if org.abbr in full_text:
            abbr_pos = full_text.find(org.abbr)
            full_pos = full_text.find(org.full)
            if full_pos < 0 or full_pos > abbr_pos:
                report.add("org", "warning",
                           f'机构名称：使用了简称"{org.abbr}"但全称'
                           f'"{org.full}"未在其前出现')


# ── 自动编号 → 纯文本（保留原有实现）────────────────────────────────────────

def _fix_auto_numbering(doc: Document, report: FormatReport) -> None:
    num_map = _build_num_map(doc)
    counters: dict[tuple, int] = {}
    missing_def_count = 0
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    for para in all_paras:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue
        numId_el = numPr.find(qn("w:numId"))
        ilvl_el  = numPr.find(qn("w:ilvl"))
        if numId_el is None:
            continue
        numId = int(numId_el.get(qn("w:val"), "0"))
        ilvl  = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
        if numId == 0:
            pPr.remove(numPr)
            continue
        abs_id, lvl_info = num_map.get(numId, (None, {}))
        info = lvl_info.get(ilvl, {})
        key = (abs_id, ilvl) if abs_id is not None else ("missing", numId, ilvl)
        counters[key] = counters.get(key, 0) + 1
        num_text = _format_num_text(info, counters[key], ilvl)
        if num_text:
            if abs_id is None:
                missing_def_count += 1
            _prepend_text_run(para, num_text)
            report.add("text", "fix",
                       f'自动编号→纯文本："{para.text[:30].strip()}" 编号 "{num_text}"')
        pPr.remove(numPr)
    if missing_def_count:
        report.add("text", "warning",
                   f"{missing_def_count} 处自动编号定义缺失，已按十进制重建，请核对编号格式")


def _build_num_map(doc: Document) -> dict:
    part = doc.part
    try:
        numbering_part = part.numbering_part
    except (AttributeError, NotImplementedError, KeyError):
        return {}
    if numbering_part is None:
        return {}
    root = numbering_part._element
    abstract_map: dict[int, dict] = {}
    for abstractNum in root.findall(qn("w:abstractNum")):
        abs_id = int(abstractNum.get(qn("w:abstractNumId"), "-1"))
        levels: dict[int, dict] = {}
        for lvl in abstractNum.findall(qn("w:lvl")):
            ilvl = int(lvl.get(qn("w:ilvl"), "0"))
            numFmt_el  = lvl.find(qn("w:numFmt"))
            lvlText_el = lvl.find(qn("w:lvlText"))
            start_el   = lvl.find(qn("w:start"))
            levels[ilvl] = {
                "numFmt":  numFmt_el.get(qn("w:val"),  "decimal") if numFmt_el  is not None else "decimal",
                "lvlText": lvlText_el.get(qn("w:val"), "%1.")     if lvlText_el is not None else "%1.",
                "start":   int(start_el.get(qn("w:val"), "1"))    if start_el   is not None else 1,
            }
        abstract_map[abs_id] = levels
    num_map: dict = {}
    for num_el in root.findall(qn("w:num")):
        numId = int(num_el.get(qn("w:numId"), "0"))
        abs_ref = num_el.find(qn("w:abstractNumId"))
        if abs_ref is None:
            continue
        abs_id = int(abs_ref.get(qn("w:val"), "-1"))
        num_map[numId] = (abs_id, abstract_map.get(abs_id, {}))
    return num_map


def _format_num_text(info: dict, count: int, ilvl: int) -> str:
    if not info:
        return f"{count}."
    num_fmt  = info.get("numFmt",  "decimal")
    lvl_text = info.get("lvlText", "%1.")
    start    = info.get("start",   1)
    n = count - 1 + start
    if num_fmt == "decimal":
        val = str(n)
    elif num_fmt == "chineseCounting":
        val = _to_chinese_num(n)
    elif num_fmt == "upperLetter":
        val = chr(ord("A") + n - 1) if 1 <= n <= 26 else str(n)
    elif num_fmt == "lowerLetter":
        val = chr(ord("a") + n - 1) if 1 <= n <= 26 else str(n)
    elif num_fmt == "bullet":
        return lvl_text if lvl_text and lvl_text != "%1" else "•"
    else:
        val = str(n)
    return re.sub(r"%\d+", val, lvl_text)


def _to_chinese_num(n: int) -> str:
    ones = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九"]
    if 1 <= n <= 9:
        return ones[n]
    if 10 <= n <= 19:
        return "十" + ones[n - 10]
    if 20 <= n <= 99:
        return ones[n // 10] + "十" + (ones[n % 10] if n % 10 else "")
    return str(n)


def _prepend_text_run(para: Any, text: str) -> None:
    p = para._p
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    first_r = p.find(qn("w:r"))
    if first_r is not None:
        p.insert(list(p).index(first_r), r)
    else:
        p.append(r)
