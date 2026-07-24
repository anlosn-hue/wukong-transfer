# word_format/body_format.py
"""
body_format.py
正文格式标准化。所有参数从 BodyConfig 读取，不再使用模块常量。
"""
from __future__ import annotations

import re
import unicodedata
from typing import Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from word_format.config_types import BodyConfig, FontConfig, TableConfig, HeadingConfig, PageConfig
from word_format.report_types import FormatReport

_SENTENCE_END = "。！？"
_PUNCT = set("。！？，、；：,.;:!?…—－）)》」』】")
_HEADING_MAX = 40
_SUB_SEP = "，：；"   # 句号太远时的次级切分标点（逗号/冒号/分号）
_CM_TO_PT = 28.3465
_TITLE_SAFETY = 0.9


def apply_body_format(doc: Document, config: BodyConfig, report: FormatReport,
                      line_pitch_pt: float = 28.95, format_tables: bool = False,
                      page_config: PageConfig | None = None,
                      wrap_title: bool = True) -> None:
    heading_patterns = _compile_heading_patterns(config.headings)
    paragraphs = doc.paragraphs
    first_heading_idx = _find_first_heading(paragraphs, heading_patterns)
    line_pitch_twip = round(line_pitch_pt * 20)  # 来自 PageConfig.line_pitch_pt，与页面网格一致

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if i < first_heading_idx:
            if i == 0:
                _fmt_title(para, config.main_title, line_pitch_twip, indent=0,
                           align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                if page_config is not None and wrap_title:
                    _apply_title_wrap(para, config.main_title, page_config)
            elif not text:
                continue
            else:
                cls = _classify_title_block_para(text)
                if cls == "salutation":
                    _fmt_title(para, config.body, line_pitch_twip, indent=0,
                               align=WD_ALIGN_PARAGRAPH.LEFT)
                elif cls == "subtitle":
                    _fmt_title(para, config.sub_title, line_pitch_twip, indent=0,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
                else:  # body
                    _fmt_title(para, config.body, line_pitch_twip, indent=2,
                               align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        else:
            level = _detect_level(text, heading_patterns)
            if level > 0:
                hcfg = next(h for h in config.headings if h.level == level)
                boundary = _heading_body_boundary(text)
                if boundary is None:
                    _fmt_title(para, FontConfig(hcfg.font, hcfg.size_pt), line_pitch_twip,
                               indent=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=hcfg.bold)
                else:
                    _apply_mixed_heading(para, text, boundary,
                                         FontConfig(hcfg.font, hcfg.size_pt), hcfg.bold,
                                         config.body, line_pitch_twip)
            else:
                _fmt_title(para, config.body, line_pitch_twip,
                           indent=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    _format_signoff(paragraphs, heading_patterns, config.body, line_pitch_twip)

    table_count = 0
    if format_tables:
        for table in doc.tables:
            _fmt_table(table, config.table)
            table_count += 1

    report.add("body", "applied",
               f"已格式化 {len(paragraphs)} 段，{table_count} 张表格")


def _compile_heading_patterns(headings: list[HeadingConfig]) -> list[tuple[int, re.Pattern]]:
    return [(h.level, re.compile(h.pattern)) for h in headings]


def _detect_level(text: str, heading_patterns: list[tuple[int, re.Pattern]]) -> int:
    for level, pattern in heading_patterns:
        if pattern.match(text):
            return level
    return 0


def _classify_title_block_para(text: str) -> str:
    """标题块内（主标题之后、首个一级标题之前）的段落归类。"""
    t = text.strip()
    if not t:
        return "body"
    last = t[-1]
    if last in "：:":
        return "salutation"
    if last not in _PUNCT:
        return "subtitle"
    return "body"


def _heading_body_boundary(text: str) -> int | None:
    """命中标题标记的段落，返回小标题/正文边界字符位置；None 表示纯标题。

    规则：段落里有**句号/分号/感叹/问号**才算混排（含正文完整句的信号；
    纯短语标题没有这些）。否则 -> None（纯标题）。混排时边界取
    「冒号/句号/分号/感叹/问号中最靠前者」（体现「标题：正文」结构）；
    若该边界在 40 字以外，改取逗号/冒号/分号中最靠前者；边界标点在末尾、
    后无正文则仍判纯标题。
    """
    t = text.strip()
    if not t:
        return None
    if not any(ch in t for ch in "。；！？"):
        return None
    cands = [t.find(ch) for ch in "：。；！？"]
    cands = [x for x in cands if x != -1]
    b = min(cands)
    if b + 1 >= len(t):
        return None
    if (b + 1) <= _HEADING_MAX:
        return b + 1
    # 边界在 40 字外：取逗号/冒号/分号中最靠前者切分
    seps = [t.find(s) for s in _SUB_SEP]
    seps = [x for x in seps if x != -1]
    if seps:
        return min(seps) + 1
    return b + 1


# 年/月/日各段允许阿拉伯数字、中文数字混用；字母 O/o 常被用来代替〇
_YMD_DIGIT = r"0-9〇○Oo零一二三四五六七八九十两壹贰叁肆伍陆柒捌玖拾"
_DATE_PATTERNS = [
    re.compile(rf"^[{_YMD_DIGIT}]{{1,4}}\s*年[{_YMD_DIGIT}]{{1,3}}\s*月"
               rf"[{_YMD_DIGIT}]{{1,3}}\s*日$"),
    re.compile(r"^\d{1,2}\s*月\d{1,2}\s*日$"),
    re.compile(r"^\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2}$"),
]


def _is_signoff_date(text: str) -> bool:
    t = text.strip()
    return any(p.match(t) for p in _DATE_PATTERNS)


_CN_DIGIT_MAP = {"〇": 0, "○": 0, "O": 0, "o": 0, "零": 0,
                 "一": 1, "二": 2, "两": 2, "三": 3, "四": 4,
                 "五": 5, "六": 6, "七": 7, "八": 8, "九": 9,
                 "壹": 1, "贰": 2, "叁": 3, "肆": 4, "伍": 5,
                 "陆": 6, "柒": 7, "捌": 8, "玖": 9, "拾": 10}


def _cn_to_int(s: str) -> int:
    """把一段年/月/日数字转为 int，支持阿拉伯、大小写中文数字、字母 O 及十位写法。"""
    s = s.strip()
    if s.isdigit():
        return int(s)
    ten = "十" if "十" in s else ("拾" if "拾" in s else None)
    if ten:
        left, _, right = s.partition(ten)
        tens = _CN_DIGIT_MAP.get(left, 1) if left else 1
        ones = _CN_DIGIT_MAP.get(right, 0) if right else 0
        return tens * 10 + ones
    digits = "".join(str(_CN_DIGIT_MAP[c]) for c in s if c in _CN_DIGIT_MAP)
    return int(digits) if digits else 0


def _normalize_date(text: str) -> str:
    """把任意识别到的日期统一改写为「YYYY年M月D日」（无年则「M月D日」）。"""
    t = text.strip()
    m = re.match(rf"^([{_YMD_DIGIT}]{{1,4}})\s*年\s*([{_YMD_DIGIT}]{{1,3}})\s*月"
                 rf"\s*([{_YMD_DIGIT}]{{1,3}})\s*日$", t)
    if m:
        y, mo, da = (_cn_to_int(x) for x in m.groups())
        return f"{y}年{mo}月{da}日"
    m = re.match(rf"^([{_YMD_DIGIT}]{{1,3}})\s*月\s*([{_YMD_DIGIT}]{{1,3}})\s*日$", t)
    if m:
        mo, da = (_cn_to_int(x) for x in m.groups())
        return f"{mo}月{da}日"
    m = re.match(r"^(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})$", t)
    if m:
        y, mo, da = (int(x) for x in m.groups())
        return f"{y}年{mo}月{da}日"
    return t


def _char_width_pt(ch: str, font_pt: float) -> float:
    if unicodedata.east_asian_width(ch) in ("W", "F", "A"):
        return font_pt
    return font_pt / 2


def _wrap_title_lines(text: str, font_pt: float, usable_pt: float) -> list[str]:
    """按 jieba 词边界把主标题折成多行。

    每行宽度不超过 usable_pt，且不超过上一行宽度——逐行收紧形成
    「上长下短」的公文梯形（第一行最长、依次递减）。
    """
    import jieba
    words = jieba.lcut(text.strip())
    lines: list[str] = []
    cur = ""
    cur_w = 0.0
    limit = usable_pt            # 第一行上限 = 版心宽；之后 < 上一行宽度
    for w in words:
        ww = sum(_char_width_pt(c, font_pt) for c in w)
        if cur and cur_w + ww > limit:
            lines.append(cur)
            limit = cur_w - 1.0   # 下一行严格短于当前行（避免相邻等长）
            cur, cur_w = w, ww
        else:
            cur += w
            cur_w += ww
    if cur:
        lines.append(cur)
    return lines or [text.strip()]


def _has_manual_break(para) -> bool:
    return bool(para._p.findall(".//" + qn("w:br")))


def _apply_title_wrap(para, main_cfg: FontConfig, page: PageConfig) -> bool:
    if _has_manual_break(para):
        return False
    usable_pt = ((page.page_width_cm - page.margin_left_cm - page.margin_right_cm)
                 * _CM_TO_PT * _TITLE_SAFETY)
    lines = _wrap_title_lines(para.text, main_cfg.size_pt, usable_pt)
    if len(lines) <= 1:
        return False
    align = para.alignment
    para.clear()
    run = para.add_run()
    for k, line in enumerate(lines):
        if k > 0:
            run.add_break(WD_BREAK.LINE)
        run.add_text(line)
    _apply_run_font(run, main_cfg.font, main_cfg.size_pt, True)
    para.alignment = align
    return True


def _find_first_heading(paragraphs: list, heading_patterns: list[tuple[int, re.Pattern]]) -> int:
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if _detect_level(text, heading_patterns) > 0:
            return i
    return 1


def _fmt_title(para, font_cfg: FontConfig, line_pitch_twip: int,
               indent: int, align: WD_ALIGN_PARAGRAPH, bold: bool = False) -> None:
    _set_para_spacing(para, align=align, indent_chars=indent,
                      line_pitch_twip=line_pitch_twip)
    _set_para_font(para, zh_font=font_cfg.font, size_pt=font_cfg.size_pt, bold=bold)


def _apply_run_font(run, zh_font: str, size_pt: float, bold: bool) -> None:
    rPr = run._r.get_or_add_rPr()
    _write_rPr_font(rPr, zh_font, str(round(size_pt * 2)), bold)


def _apply_mixed_heading(para, text: str, boundary: int,
                         head_cfg: FontConfig, head_bold: bool,
                         body_cfg: FontConfig, line_pitch_twip: int) -> None:
    head = text[:boundary]
    tail = text[boundary:]
    # 段落级：标准正文（首行缩进2字、两端对齐）
    _set_para_spacing(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      indent_chars=2, line_pitch_twip=line_pitch_twip)
    para.clear()
    r1 = para.add_run(head)
    _apply_run_font(r1, head_cfg.font, head_cfg.size_pt, head_bold)
    r2 = para.add_run(tail)
    _apply_run_font(r2, body_cfg.font, body_cfg.size_pt, False)
    # 段落标记字体设为正文
    para_rPr = _get_or_create(para._p.get_or_add_pPr(), "w:rPr")
    _write_rPr_font(para_rPr, body_cfg.font, str(round(body_cfg.size_pt * 2)), False)


def _set_right_indent_chars(para, end_chars: float) -> None:
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    # 清掉首行缩进，避免与右缩进冲突
    for a in ("w:firstLineChars", "w:firstLine"):
        ind.attrib.pop(qn(a), None)
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:endChars"), str(round(end_chars * 100)))   # 单位 1/100 字，支持半字


def _set_space_before(para, before_twip: int) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before_twip))
    spacing.attrib.pop(qn("w:beforeLines"), None)


def _format_signoff(paragraphs, heading_patterns, body_cfg: FontConfig,
                    line_pitch_twip: int) -> bool:
    # 找最后一个非空段
    idx = len(paragraphs) - 1
    while idx >= 0 and not paragraphs[idx].text.strip():
        idx -= 1
    if idx < 0:
        return False
    date_para = paragraphs[idx]
    if not _is_signoff_date(date_para.text):
        return False
    # 统一改写为标准格式「YYYY年M月D日」
    std = _normalize_date(date_para.text)
    if std != date_para.text.strip():
        date_para.clear()
        r = date_para.add_run(std)
        _apply_run_font(r, body_cfg.font, body_cfg.size_pt, False)
    date_len = len(std)
    # 往上找紧邻非空段作落款
    j = idx - 1
    while j >= 0 and not paragraphs[j].text.strip():
        j -= 1
    sign_para = None
    sign_len = date_len
    if j >= 0:
        cand = paragraphs[j].text.strip()
        if (len(cand) <= 30 and cand[-1] not in _SENTENCE_END
                and _detect_level(cand, heading_patterns) == 0
                and not _is_signoff_date(cand)):
            sign_para = paragraphs[j]
            sign_len = len(cand)
    if sign_para is not None:
        _set_right_indent_chars(sign_para, 2)
    # 日期中心对齐落款中心（浮点，支持半字精度）
    date_indent = max(0.0, 2 + (sign_len - date_len) / 2.0)
    _set_right_indent_chars(date_para, date_indent)
    # 落款块第一段（有落款则落款，否则日期）与前面正文空两行
    head_para = sign_para if sign_para is not None else date_para
    _set_space_before(head_para, 2 * line_pitch_twip)
    return True


def _fmt_table(table, config: TableConfig) -> None:
    tbl = table._tbl
    tblPr = _get_or_create(tbl, "w:tblPr")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "autofit")

    table_twip = round(config.line_spacing_pt * 20)
    for row in table.rows:
        trPr = _get_or_create(row._tr, "w:trPr")
        trHeight = trPr.find(qn("w:trHeight"))
        if trHeight is None:
            trHeight = OxmlElement("w:trHeight")
            trPr.append(trHeight)
        trHeight.set(qn("w:hRule"), "auto")
        for cell in row.cells:
            for para in cell.paragraphs:
                _set_table_para_spacing(para, table_twip)
                _set_para_font(para, zh_font=config.font, size_pt=config.size_pt, bold=False)


def _set_para_spacing(para, align, indent_chars: int, line_pitch_twip: int) -> None:
    para.alignment = align
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"),     str(line_pitch_twip))
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:before"),   "0")
    spacing.set(qn("w:after"),    "0")
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if indent_chars > 0:
        ind.set(qn("w:firstLineChars"), str(indent_chars * 100))
        ind.attrib.pop(qn("w:firstLine"), None)
    else:
        ind.attrib.pop(qn("w:firstLineChars"), None)
        ind.attrib.pop(qn("w:firstLine"), None)
        ind.set(qn("w:firstLine"), "0")


def _set_table_para_spacing(para, twip: int) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"),     str(twip))
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:before"),   "0")
    spacing.set(qn("w:after"),    "0")
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        pPr.remove(ind)


def _set_para_font(para, zh_font: str, size_pt: float, bold: bool) -> None:
    half_pts = str(round(size_pt * 2))
    pPr = para._p.get_or_add_pPr()
    para_rPr = _get_or_create(pPr, "w:rPr")
    _write_rPr_font(para_rPr, zh_font, half_pts, bold)
    for run in para.runs:
        rPr = run._r.get_or_add_rPr()
        _write_rPr_font(rPr, zh_font, half_pts, bold)


def _write_rPr_font(rPr, zh_font: str, half_pts: str, bold: bool) -> None:
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    zh_font)
    rFonts.set(qn("w:hAnsi"),    zh_font)
    rFonts.set(qn("w:eastAsia"), zh_font)
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        rFonts.attrib.pop(qn(attr), None)
    for ptag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(ptag))
        if el is None:
            el = OxmlElement(ptag)
            rPr.append(el)
        el.set(qn("w:val"), half_pts)
    b_el   = rPr.find(qn("w:b"))
    bCs_el = rPr.find(qn("w:bCs"))
    if bold:
        if b_el   is None: rPr.append(OxmlElement("w:b"))
        if bCs_el is None: rPr.append(OxmlElement("w:bCs"))
    else:
        if b_el   is not None: rPr.remove(b_el)
        if bCs_el is not None: rPr.remove(bCs_el)


def _get_or_create(parent, prefixed: str) -> Any:
    el = parent.find(qn(prefixed))
    if el is None:
        el = OxmlElement(prefixed)
        parent.append(el)
    return el
